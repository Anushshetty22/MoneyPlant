from pathlib import Path
import re
import argparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "data-source-decision.md"
OUTPUT = ROOT / "docs" / "MoneyPlant_Phase1_Data_Sources.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"


def set_run_font(run, name="Arial", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "C7D2E0")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text, size=11):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(https?://[^)]+\)|https?://\S+)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[(.*?)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("http"):
            add_hyperlink(paragraph, token.rstrip(".,"), token.rstrip(".,"))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Roboto Mono", size=max(9, size - 1), color=INK)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size)


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body(doc, text, style=None, before=0, after=6, size=11):
    paragraph = doc.add_paragraph(style=style)
    style_paragraph(paragraph, before=before, after=after)
    add_inline(paragraph, text, size=size)
    return paragraph


def add_table(doc, rows):
    headers = rows[0]
    cols = len(headers)
    if cols == 4:
        widths = [1500, 1900, 3000, 2960]
    elif cols == 5:
        # Database-design tables use a compact key/type/null/default layout
        # and reserve most of the width for the explanation column.
        widths = [1500, 1800, 700, 1700, 3660]
    elif cols == 3:
        widths = [2300, 2200, 4860]
    else:
        widths = [3120, 6240]
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.15)
        run = p.add_run(value)
        set_run_font(run, size=9.5, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            style_paragraph(p, after=0, line=1.15)
            add_inline(p, value, size=9.5)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=5, line=1.0)
    return table


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-+:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build(source=SOURCE, output=OUTPUT, title="Phase 1 Data Sources", subtitle="Collection plan, field definitions, and pre-schema checklist"):
    source_lines = Path(source).read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 18, BLACK, 18, 8),
        ("Heading 2", 14, BLACK, 14, 6),
        ("Heading 3", 12, BLACK, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    # Simple title block matching the supplied MoneyPlant brief.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=4, after=2, line=1.0)
    run = p.add_run("MoneyPlant")
    set_run_font(run, size=26, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=0, after=2, line=1.0)
    run = p.add_run(title)
    set_run_font(run, size=18, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=0, after=18, line=1.15)
    run = p.add_run(subtitle)
    set_run_font(run, size=11, color=BLACK, italic=True)

    skip_h1 = True
    index = 0
    in_code = False
    code_lines = []
    current_decimal_num_id = None
    while index < len(source_lines):
        line = source_lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                style_paragraph(p, before=2, after=8, line=1.0)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Roboto Mono", size=9, color=INK)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if stripped == "":
            current_decimal_num_id = None
            index += 1
            continue
        if stripped.startswith("# ") and skip_h1:
            skip_h1 = False
            index += 1
            continue
        if stripped.startswith("Status:"):
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(source_lines) and "|---" in source_lines[index + 1]:
            rows, index = parse_table(source_lines, index)
            add_table(doc, rows)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading_match:
            current_decimal_num_id = None
            level = len(heading_match.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading_match.group(2), size={1: 16, 2: 13, 3: 12}[level])
            index += 1
            continue
        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet_match or number_match:
            text = (bullet_match or number_match).group(1)
            if number_match:
                number_value = int(re.match(r"^(\d+)", stripped).group(1))
                if current_decimal_num_id is None or number_value == 1:
                    current_decimal_num_id = create_decimal_numbering(doc)
                p = doc.add_paragraph()
                apply_numbering(p, current_decimal_num_id)
            else:
                current_decimal_num_id = None
                p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            style_paragraph(p, after=2, line=1.15)
            add_inline(p, text, size=11)
            index += 1
            continue
        add_body(doc, stripped)
        index += 1

    # Ensure list styles have explicit typography and spacing.
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.15

    doc.core_properties.title = f"MoneyPlant {title}"
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "MoneyPlant"
    doc.core_properties.keywords = "MoneyPlant, data engineering, data sources, PostgreSQL"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", default=str(SOURCE))
    parser.add_argument("--output", default=str(OUTPUT))
    parser.add_argument("--title", default="Phase 1 Data Sources")
    parser.add_argument("--subtitle", default="Collection plan, field definitions, and pre-schema checklist")
    args = parser.parse_args()
    build(args.source, args.output, args.title, args.subtitle)
